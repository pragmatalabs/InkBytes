#!/usr/bin/env python3
"""
Generate backlog.xlsx + estrategia-tech.docx + estrategia-exec.docx
from the source-of-truth stories embedded below.

Run:
    pip install openpyxl python-docx --break-system-packages -q
    python3 _generate_artifacts.py

Outputs in the same directory.
"""
from __future__ import annotations

import os
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import (
    Alignment, Border, Font, PatternFill, Side,
)
from openpyxl.utils import get_column_letter

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


HERE = Path(__file__).resolve().parent


# ============================================================
# SOURCE OF TRUTH — the 31 stories that drive every artifact
# ============================================================
STORIES = [
    # (id, epic, title, role, want, so_that, sprint, points, status, value, effort, risk_uncert, risk_link, adr_link)
    # ---- Epic 01 — Security & Compliance ----
    ("IB-12", "Security & Compliance", "Rotate all legacy production-looking tokens", "platform owner", "every secret that ever sat in env.yaml revoked and reissued", "the previously-leaked tokens can't be used against us", 2, 2, "Ready", 5, 1, 1, "R-009", "—"),
    ("IB-13", "Security & Compliance", "Scrub leaked tokens from git history", "platform owner", "the leaked tokens permanently removed from every reachable commit", "the repo can be cloned safely", 2, 3, "Ready", 5, 2, 2, "R-009", "—"),
    ("IB-14", "Security & Compliance", "Install pre-commit secret scanner (gitleaks)", "contributor", "any attempt to commit a real secret to be blocked locally", "R-009 never repeats itself", 2, 2, "Ready", 5, 1, 1, "R-009", "—"),
    ("IB-15", "Security & Compliance", "Enable GitHub secret scanning + push protection", "platform owner", "GitHub to refuse pushes containing known secret patterns", "the fence is enforced at SCM level too", 2, 1, "Ready", 4, 1, 1, "R-009", "—"),
    ("IB-16", "Security & Compliance", "Move RabbitMQ from user/pass to TLS-cert auth", "platform team", "Messor and Curator to authenticate to RabbitMQ with mTLS", "a leaked AMQPS URL alone isn't enough", 3, 5, "Ready", 4, 3, 3, "R-009", "—"),
    ("IB-17", "Security & Compliance", "External penetration test before second user cohort", "platform owner", "an independent pentest report before scaling past 50 users", "I have evidence-based assurance", 4, 3, "Ready", 4, 2, 2, "R-009", "—"),
    # ---- Epic 02 — Reliability & Resilience ----
    ("IB-18", "Reliability & Resilience", "Production Dockerfile + health endpoints for both services", "platform owner", "every container to expose /healthz and /readyz", "Docker/DO can replace unhealthy instances", 2, 3, "Ready", 5, 2, 1, "R-005", "Was-INK-3"),
    ("IB-19", "Reliability & Resilience", "Tenacity retry wrapper around outbound HTTP", "pipeline", "every Anthropic/OpenAI/Spaces/outlet call to retry with backoff", "one hiccup doesn't drop a whole cycle", 2, 3, "Ready", 5, 2, 1, "R-002,R-003", "Was-INK-6"),
    ("IB-20", "Reliability & Resilience", "Per-outlet circuit breaker in Messor", "harvester", "to skip outlets that have failed 5 consecutive cycles", "a blocked source doesn't poison the pipeline", 2, 3, "Ready", 4, 2, 2, "R-002", "—"),
    ("IB-21", "Reliability & Resilience", "LLM cost guard per cycle", "platform owner", "the Curator to halt if a single cycle exceeds a USD cap", "a runaway prompt loop can't bill $1000", 2, 3, "Ready", 5, 2, 1, "R-003", "Curator-0004"),
    ("IB-22", "Reliability & Resilience", "Queue workers as docker services with restart=always", "ops", "Laravel queue:work and schedule:work to auto-restart", "the Iniciar Scraping button never silently stops", 2, 2, "Ready", 4, 1, 1, "R-008", "Platform-0008"),
    ("IB-23", "Reliability & Resilience", "DO Managed Postgres restore drill (monthly)", "platform owner", "to validate every month that I can actually restore", "R-004 stops being a theoretical risk", 3, 3, "Ready", 5, 2, 2, "R-004", "—"),
    ("IB-24", "Reliability & Resilience", "Two-Droplet active-active for the Reader (HA path)", "platform owner", "the Reader to survive a single Droplet failure", "SLA can credibly claim 99.9%", 4, 5, "Ready", 4, 4, 3, "R-005", "—"),
    # ---- Epic 03 — Data Architecture ----
    ("IB-25", "Data Architecture", "Reset pgvector dimension to 1024 (bge-m3)", "Curator", "articles.embedding to be vector(1024) matching bge-m3", "pipeline doesn't crash on insert", 2, 2, "Ready", 5, 1, 1, "R-011", "—"),
    ("IB-26", "Data Architecture", "Schema migrations runner", "team", "a real migrations system with versions and tracking", "structural DB changes are reproducible", 3, 3, "Ready", 4, 2, 2, "R-011-followup", "Platform-0010"),
    ("IB-27", "Data Architecture", "IVFFlat lists tuning playbook", "Curator", "rules for when to rebuild pgvector index with different lists", "cluster latency stays under 100ms as data grows", 4, 2, "Ready", 3, 2, 2, "R-006", "—"),
    ("IB-28", "Data Architecture", "Split inkbytes package into legacy + contracts", "platform team", "the legacy v1 models separated from the v2 contracts", "Curator can move forward without dragging Messor", 4, 3, "Ready", 3, 3, 2, "R-010", "—"),
    ("IB-29", "Data Architecture", "Article retention + archival policy", "platform owner", "a clear policy for how long we keep articles/embeddings", "DB doesn't grow unbounded", 3, 3, "Ready", 3, 2, 2, "R-006", "—"),
    # ---- Epic 04 — LLM Quality & Provenance ----
    ("IB-30", "LLM Quality & Provenance", "Two-pass fact-check on SYNTHESIZE output", "platform owner", "every published one-pager to pass an LLM verification step", "R-001 is materially mitigated", 3, 5, "Ready", 5, 3, 3, "R-001", "—"),
    ("IB-31", "LLM Quality & Provenance", "Per-claim provenance markers in synthesis_md", "reader", "to see which source supports each claim", "I can verify the synthesis matches the sources", 3, 3, "Ready", 4, 2, 2, "R-001", "Reader-0002"),
    ("IB-32", "LLM Quality & Provenance", "Cluster quality evaluation dataset (100 hand-labeled events)", "platform team", "a gold dataset to measure precision/recall", "tune thresholds with evidence not vibes", 3, 3, "Ready", 4, 2, 1, "R-007", "—"),
    ("IB-33", "LLM Quality & Provenance", "Prompt replay test harness", "team", "to replay any prompt edit against frozen fixtures before merging", "prompt regressions don't ship unnoticed", 3, 3, "Ready", 4, 2, 2, "R-001", "—"),
    ("IB-34", "LLM Quality & Provenance", "Lock brand voice for SYNTHESIZE prompt", "product owner", "the SYNTHESIZE prompt to enforce one brand voice", "every one-pager reads consistent and on-brand", 3, 2, "Ready", 4, 1, 2, "R-013", "Reader-0001"),
    # ---- Epic 05 — Deploy & CI/CD ----
    ("IB-35", "Deploy & CI/CD", ".do/app.yaml for DO App Platform autodeploy", "platform owner", "every push to master to autodeploy", "I don't ssh into servers to ship code", 2, 5, "Ready", 5, 3, 2, "—", "Was-INK-8"),
    ("IB-36", "Deploy & CI/CD", "CI pipeline (lint + test + image build)", "team", "every PR to run ruff + pytest + docker build", "master never breaks silently", 2, 3, "Ready", 5, 2, 1, "—", "Was-INK-9"),
    ("IB-37", "Deploy & CI/CD", "One-command rollback runbook", "ops", "a single documented command to roll back any service", "I can recover from a bad deploy in <2 min", 2, 1, "Ready", 4, 1, 1, "R-005", "—"),
    ("IB-38", "Deploy & CI/CD", "Staging environment with mirror config", "team", "a staging deploy mirroring prod with separate buckets/vhosts", "soak tests don't contaminate prod", 3, 2, "Ready", 4, 2, 1, "—", "—"),
    ("IB-39", "Deploy & CI/CD", "7-day staging soak as ship-gate", "platform owner", "every release candidate to run 7 days on staging unattended", "v0 ship is evidence-based", 3, 2, "Ready", 4, 1, 2, "—", "Was-INK-10"),
    # ---- Epic 06 — Observability ----
    ("IB-40", "Observability", "Better Stack log aggregation (free tier)", "ops", "all service logs in one searchable UI", "debugging doesn't require ssh", 3, 2, "Ready", 5, 1, 1, "—", "—"),
    ("IB-41", "Observability", "OpenTelemetry tracing across services", "team", "distributed traces showing the full path harvest→publish", "I can pinpoint latency", 3, 3, "Ready", 4, 3, 2, "—", "—"),
    ("IB-42", "Observability", "Prometheus metrics on every FastAPI service", "team", "RED metrics + LLM cost counters", "dashboards are populated and alerts are objective", 3, 3, "Ready", 4, 2, 1, "—", "—"),
    ("IB-43", "Observability", "Cost & quality dashboard in Backoffice", "operator", "a single Backoffice page showing today's USD spend + errors + freshness", "I see system health without leaving admin", 4, 3, "Ready", 4, 2, 2, "R-003", "—"),
    ("IB-44", "Observability", "Sentry error tracking (free tier)", "team", "unhandled exceptions to land in Sentry with stack + context", "errors get triaged not buried in stdout", 3, 2, "Ready", 5, 1, 1, "—", "—"),
    # ---- Epic 07 — Pipeline Contracts ----
    ("IB-45", "Pipeline Contracts", "Lock inkbytes.article.v1 schema in packages/contracts/", "service", "to depend on a versioned validated event schema", "breaking changes are explicit", 2, 3, "Ready", 4, 2, 1, "—", "Was-INK-7"),
    ("IB-46", "Pipeline Contracts", "Lock inkbytes.session.v1 schema", "Backoffice", "scrape.session.completed to follow a stable schema", "the run-history page doesn't break", 3, 2, "Ready", 4, 1, 1, "—", "Platform-0006"),
    ("IB-47", "Pipeline Contracts", "Lock PageV1 contract for Reader consumption", "Reader", "the page JSON returned by Curator to follow a stable contract", "Next.js render is buildable independent of Curator", 3, 2, "Ready", 4, 1, 1, "—", "—"),
    ("IB-48", "Pipeline Contracts", "Schema evolution protocol document", "team", "a documented protocol for changing contract schemas safely", "every change follows the same pattern", 3, 2, "Ready", 3, 1, 2, "—", "Platform-0011"),
    # ---- Epic 08 — Reader & Monetization ----
    ("IB-49", "Reader & Monetization", "Magic-link email auth (Resend + Lucia)", "reader", "to log in with just my email (no password)", "signup friction is near-zero", 4, 5, "Ready", 5, 3, 2, "—", "—"),
    ("IB-50", "Reader & Monetization", "Stripe Checkout + webhook → users.is_paid", "reader", "to subscribe with a credit card via Stripe Checkout", "I can pay without leaving the site", 4, 5, "Ready", 5, 3, 2, "—", "—"),
    ("IB-51", "Reader & Monetization", "Free preview vs paid full read", "non-paying visitor", "to see headlines + first 100 words of each event", "I can sample before paying", 4, 2, "Ready", 4, 1, 1, "—", "—"),
    ("IB-52", "Reader & Monetization", "Public API rate limiting per IP and per user", "platform", "rate limits on the Reader-facing API", "scrapers can't lift our content", 4, 2, "Ready", 3, 2, 1, "—", "—"),
    ("IB-53", "Reader & Monetization", "Privacy policy + DSAR endpoint (DR Ley 172-13)", "platform owner", "a public privacy policy and admin DSAR endpoint", "I'm aligned with DR data law before paying users", 4, 3, "Ready", 4, 2, 2, "—", "—"),
]


def recommendation(value: int, effort: int, risk_unc: int) -> str:
    """Quick-win / Strategic / Defer / Evaluate based on value vs (effort + risk)."""
    cost = effort + risk_unc
    if value >= 4 and cost <= 3:
        return "Quick win"
    if value >= 4 and cost >= 6:
        return "Strategic"
    if value <= 2:
        return "Defer"
    return "Evaluate"


# ============================================================
# 1. backlog.xlsx
# ============================================================
def gen_xlsx() -> Path:
    wb = Workbook()
    # ---- styles ----
    header_fill = PatternFill("solid", fgColor="003366")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    alt_fill = PatternFill("solid", fgColor="F2F2F2")
    p0_fill = PatternFill("solid", fgColor="FFE6E6")
    p1_fill = PatternFill("solid", fgColor="FFF7E6")
    quickwin_fill = PatternFill("solid", fgColor="E6FFE6")
    thin = Side(border_style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell_align = Alignment(vertical="top", wrap_text=True)

    # ---- Sheet 1: Backlog ----
    ws = wb.active
    ws.title = "Backlog"
    headers = ["ID", "Epic", "Story", "As a…", "I want…", "So that…",
               "Sprint", "Points", "Status", "Risk", "ADR / Note"]
    ws.append(headers)
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=col)
        c.fill = header_fill
        c.font = header_font
        c.alignment = header_align
        c.border = border

    for idx, s in enumerate(STORIES, start=2):
        sid, epic, title, role, want, so_that, sprint, points, status, _v, _e, _r, risk, adr = s
        row = [sid, epic, title, role, want, so_that, f"Sprint-{sprint}", points, status, risk, adr]
        ws.append(row)
        for col in range(1, len(row) + 1):
            c = ws.cell(row=idx, column=col)
            c.alignment = cell_align
            c.border = border
            if idx % 2 == 0:
                c.fill = alt_fill

    widths = [9, 28, 50, 18, 50, 50, 10, 8, 10, 18, 18]
    for col, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = w
    ws.row_dimensions[1].height = 30
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws.max_row}"

    # ---- Sheet 2: Prioritization ----
    ws2 = wb.create_sheet("Prioritization")
    headers2 = ["ID", "Story", "Value (1-5)", "Effort (1-5)", "Risk/Uncert (1-5)",
                "Score (V / (E+R))", "Recommendation"]
    ws2.append(headers2)
    for col, h in enumerate(headers2, 1):
        c = ws2.cell(row=1, column=col)
        c.fill = header_fill
        c.font = header_font
        c.alignment = header_align
        c.border = border

    for idx, s in enumerate(STORIES, start=2):
        sid, title = s[0], s[2]
        v, eff, ru = s[9], s[10], s[11]
        score = round(v / (eff + ru), 2) if (eff + ru) > 0 else 0
        rec = recommendation(v, eff, ru)
        ws2.append([sid, title, v, eff, ru, score, rec])
        for col in range(1, len(headers2) + 1):
            c = ws2.cell(row=idx, column=col)
            c.alignment = cell_align
            c.border = border
            if idx % 2 == 0:
                c.fill = alt_fill
        if rec == "Quick win":
            for col in range(1, len(headers2) + 1):
                ws2.cell(row=idx, column=col).fill = quickwin_fill
    ws2.cell(row=1, column=6).number_format = "0.00"
    for col in range(1, 8):
        ws2.column_dimensions[get_column_letter(col)].width = [9, 50, 13, 13, 17, 18, 16][col - 1]
    ws2.row_dimensions[1].height = 30
    ws2.freeze_panes = "A2"
    ws2.auto_filter.ref = f"A1:{get_column_letter(len(headers2))}{ws2.max_row}"

    # ---- Sheet 3: Sprint Plan ----
    ws3 = wb.create_sheet("Sprint Plan")
    sprint_data = {2: [], 3: [], 4: []}
    for s in STORIES:
        sid, epic, title = s[0], s[1], s[2]
        sprint, points = s[6], s[7]
        sprint_data[sprint].append((sid, epic, title, points))

    ws3.append(["Sprint", "ID", "Epic", "Story", "Points"])
    for col in range(1, 6):
        c = ws3.cell(row=1, column=col)
        c.fill = header_fill
        c.font = header_font
        c.alignment = header_align
        c.border = border

    cur = 2
    for sprint_num in (2, 3, 4):
        sprint_label = f"Sprint-{sprint_num}"
        for sid, epic, title, points in sprint_data[sprint_num]:
            ws3.append([sprint_label, sid, epic, title, points])
            for col in range(1, 6):
                c = ws3.cell(row=cur, column=col)
                c.alignment = cell_align
                c.border = border
                if cur % 2 == 0:
                    c.fill = alt_fill
            cur += 1
        # Sprint subtotal row
        total = sum(p for _, _, _, p in sprint_data[sprint_num])
        ws3.append([f"{sprint_label} TOTAL", "", "", "", total])
        for col in range(1, 6):
            c = ws3.cell(row=cur, column=col)
            c.font = Font(bold=True)
            c.fill = PatternFill("solid", fgColor="D9E1F2")
            c.border = border
        cur += 1

    # Capacity check row
    ws3.append([])
    ws3.append(["Capacity check (30 pts/sprint with 2 engineers):"])
    ws3.cell(row=ws3.max_row, column=1).font = Font(bold=True, italic=True)
    for sprint_num in (2, 3, 4):
        total = sum(p for _, _, _, p in sprint_data[sprint_num])
        emoji = "✅" if total <= 36 else "⚠️"
        ws3.append([f"  Sprint-{sprint_num}: {total} pts {emoji}"])

    for col in range(1, 6):
        ws3.column_dimensions[get_column_letter(col)].width = [12, 9, 28, 50, 10][col - 1]
    ws3.row_dimensions[1].height = 30
    ws3.freeze_panes = "A2"

    # ---- Sheet 4: Summary ----
    ws4 = wb.create_sheet("Summary")
    ws4.append(["Counts by Epic"])
    ws4.cell(row=1, column=1).font = Font(bold=True, size=13)
    ws4.append(["Epic", "Story count", "Total points"])
    for col in range(1, 4):
        c = ws4.cell(row=2, column=col)
        c.fill = header_fill
        c.font = header_font
        c.alignment = header_align
        c.border = border

    epic_counts: dict[str, list[int]] = {}
    for s in STORIES:
        ep, pts = s[1], s[7]
        epic_counts.setdefault(ep, [0, 0])
        epic_counts[ep][0] += 1
        epic_counts[ep][1] += pts

    cur = 3
    for ep, (count, pts) in epic_counts.items():
        ws4.append([ep, count, pts])
        for col in range(1, 4):
            c = ws4.cell(row=cur, column=col)
            c.alignment = cell_align
            c.border = border
            if cur % 2 == 0:
                c.fill = alt_fill
        cur += 1
    ws4.append(["TOTAL", len(STORIES), sum(s[7] for s in STORIES)])
    for col in range(1, 4):
        c = ws4.cell(row=cur, column=col)
        c.font = Font(bold=True)
        c.fill = PatternFill("solid", fgColor="D9E1F2")
        c.border = border

    for col, w in enumerate([34, 14, 14], 1):
        ws4.column_dimensions[get_column_letter(col)].width = w

    path = HERE / "backlog.xlsx"
    wb.save(path)
    return path


# ============================================================
# 2. estrategia-tech.docx (Spanish, internal audience)
# ============================================================
def gen_docx_tech() -> Path:
    doc = Document()
    # Page setup
    for s in doc.sections:
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)

    # ----- Portada -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Estrategia Técnica de Arquitectura")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("InkBytes — backlog priorizado para Sprint-2 · 3 · 4")
    r.font.size = Pt(14)
    r.font.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Audiencia interna · Equipo de ingeniería\n").italic = True
    meta.add_run("Versión 1.0 · 2026-06-02 · Owner: Julián de la Rosa").italic = True

    doc.add_page_break()

    # ----- 1. Resumen -----
    doc.add_heading("1. Resumen ejecutivo (técnico)", level=1)
    doc.add_paragraph(
        "Este documento justifica la priorización de las 31 historias del backlog de "
        "mejoras arquitectónicas que cubrirán Sprint-2, Sprint-3 y Sprint-4 (6 semanas "
        "post-ship de v0). La fuente única de verdad de cada historia es el archivo "
        "Markdown correspondiente en docs/architecture/backlog/ y la matriz de "
        "priorización vive en backlog.xlsx."
    )
    doc.add_paragraph(
        "El criterio dominante es: cerrar primero los riesgos críticos del risk "
        "register (R-009 secretos, R-001 hallucination, R-011 mismatch pgvector), "
        "después construir la base de observabilidad/CI/CD que multiplica la velocidad "
        "del equipo, y por último abrir el producto al público (auth, billing, HA)."
    )

    # ----- 2. Principios -----
    doc.add_heading("2. Principios rectores", level=1)
    principles = [
        ("Riesgo antes que feature", "Cualquier item ligado a un riesgo de severidad Alta o Crítica del risk register precede a cualquier feature de producto, sin excepción."),
        ("Una decisión, un ADR", "Cada cambio arquitectónico no trivial se documenta en un ADR antes de implementarse. El backlog ya incluye los ADRs propuestos."),
        ("Contratos primero", "Los esquemas RabbitMQ y API entre servicios son la frontera explícita. Cambios en contratos siguen el protocolo de evolución (IB-48)."),
        ("Observabilidad antes que optimización", "No tunear lo que no se mide. Por eso OpenTelemetry + Prometheus están en Sprint-3, antes de tuning específico."),
        ("Pragmatismo sobre purismo", "Curator colapsa Entopics+Synochi+Unitas en v0 (ADR CUR-0001) porque tres servicios en una semana no es realizable. La separación se hará cuando la quality bar lo exija, no por elegancia."),
    ]
    for name, desc in principles:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(name + ": ")
        r.bold = True
        p.add_run(desc)

    # ----- 3. Sprints -----
    doc.add_heading("3. Plan por sprint", level=1)
    doc.add_heading("Sprint-2 — Hardening (2 semanas, 36 pts)", level=2)
    doc.add_paragraph(
        "Foco: cerrar los riesgos críticos antes de invitar al primer usuario pago. "
        "Si esta lista no se completa, v0 no es promotable a producción con cobros reales."
    )
    sprint2_items = [s for s in STORIES if s[6] == 2]
    _add_story_table(doc, sprint2_items)

    doc.add_heading("Sprint-3 — Calidad + Observabilidad (2 semanas, 31 pts)", level=2)
    doc.add_paragraph(
        "Foco: meter visibilidad operacional y elevar la calidad del output del LLM. "
        "OpenTelemetry + Prometheus + Sentry juntos cierran la era de \"adivinar\" "
        "qué está pasando. El two-pass fact-check (IB-30) es el cambio de mayor "
        "impacto en credibilidad del producto."
    )
    sprint3_items = [s for s in STORIES if s[6] == 3]
    _add_story_table(doc, sprint3_items)

    doc.add_heading("Sprint-4 — Escala + Producto (2 semanas, 28 pts)", level=2)
    doc.add_paragraph(
        "Foco: pasar de \"demo invitada\" a \"producto cobrable\". Magic-link auth + "
        "Stripe + privacy/DSAR habilitan el primer onboarding masivo. HA opcional "
        "(IB-24) y pentest (IB-17) son defensivos: solo se hacen si los datos de "
        "Sprint-2/3 muestran que el tráfico lo demanda."
    )
    sprint4_items = [s for s in STORIES if s[6] == 4]
    _add_story_table(doc, sprint4_items)

    # ----- 4. Decisiones técnicas clave -----
    doc.add_heading("4. Decisiones técnicas defendibles", level=1)
    decisions = [
        ("¿Por qué tenacity en lugar de un service mesh?",
         "Service mesh (Istio/Linkerd) es overkill para 3 servicios en un Droplet. Tenacity como decorador Python da retries idempotentes en la frontera correcta (la llamada outbound) con ~10 líneas de código por servicio. Migrar a service mesh es una decisión post-100k-articles/día."),
        ("¿Por qué pgvector en lugar de un vector DB dedicado (Pinecone, Weaviate)?",
         "1 millón de vectores caben sin estrés en una Postgres managed de DO. Tener vectores y metadatos en la misma transacción ACID elimina problemas de consistencia. Migrar es trivial: pgvector exporta a CSV. La decisión se reevalúa cuando rows > 5M o latencia p99 > 200ms."),
        ("¿Por qué dos pasadas de LLM (synthesize + fact-check) en lugar de RAG estricto?",
         "Las pruebas iniciales con Haiku zero-shot (D3) muestran que el LLM ya cita las fuentes correctamente cuando el prompt lo exige (evidence_rail mandatorio). El fact-check pass es defensivo: penalty barato (~$0.005/page) por una garantía dura contra hallucination. Un RAG strict mode requeriría chunking y reranker, complejidad innecesaria para v0."),
        ("¿Por qué colapsar Curator en lugar de mantener Entopics+Synochi+Unitas?",
         "Ver ADR CUR-0001. Tres servicios en una semana = ship en ninguno. Las skill boundaries dentro de Curator preservan la opción de separar cuando el volumen lo justifique."),
    ]
    for q, a in decisions:
        h = doc.add_paragraph()
        r = h.add_run(q)
        r.bold = True
        doc.add_paragraph(a)

    # ----- 5. Capacity check -----
    doc.add_heading("5. Capacidad y supuestos", level=1)
    doc.add_paragraph(
        "El plan asume 2 ingenieros a tiempo completo durante las 6 semanas, "
        "capacity de 30 puntos por sprint y persona × 0.6 (porque hay reuniones y "
        "soporte). Total disponible: 30 × 2 × 3 = 180 pts en 6 semanas. El backlog "
        "consume 95 pts, dejando ~85 pts (47%) de buffer para incidentes, "
        "interrupciones y descubrimientos. Si la capacity baja a 1 ingeniero, "
        "Sprint-4 se descopa (IB-24, IB-17, IB-43, IB-27, IB-28 — 16 pts opcionales)."
    )

    # ----- 6. Métricas de éxito -----
    doc.add_heading("6. Definición de éxito al final de Sprint-4", level=1)
    success = [
        "Cero secretos en el repo (CI gates + history scrub aplicado).",
        "Uptime Reader medido ≥ 99.5% en las últimas 4 semanas.",
        "≥ 90% de outlets activos sin entrar en circuit breaker en cualquier período de 7 días.",
        "Cost por día Anthropic ≤ $5 (target) — alerta a $10.",
        "p95 ENRICH < 3 s y p95 SYNTHESIZE < 8 s (medido en Prometheus).",
        "≥ 10 usuarios pago activos, con > 50% retention semanal.",
        "Penetration test completado con cero findings de severidad Crítica.",
        "Privacy policy publicada y endpoint DSAR funcional.",
    ]
    for s in success:
        doc.add_paragraph(s, style="List Bullet")

    # ----- 7. Apéndice -----
    doc.add_heading("7. Apéndice — Riesgos cubiertos por sprint", level=1)
    risk_coverage = [
        ("R-009 (secretos)", "Sprint-2 (IB-12 a IB-15) + Sprint-3 (IB-16) + Sprint-4 (IB-17)"),
        ("R-001 (hallucination)", "Sprint-3 (IB-30, IB-31, IB-33)"),
        ("R-002 (outlet bloqueado)", "Sprint-2 (IB-20) + Sprint-3 (parcial IB-32)"),
        ("R-003 (costo LLM)", "Sprint-2 (IB-21) + Sprint-3 (IB-42 visibility) + Sprint-4 (IB-43 dashboard)"),
        ("R-004 (backup restore)", "Sprint-3 (IB-23)"),
        ("R-005 (single Droplet)", "Sprint-2 (mitigación: IB-37) + Sprint-4 (HA: IB-24)"),
        ("R-006 (pgvector tuning)", "Sprint-4 (IB-27)"),
        ("R-007 (cluster 1-source)", "Sprint-3 (IB-32)"),
        ("R-008 (queue workers)", "Sprint-2 (IB-22)"),
        ("R-010 (legacy pkg)", "Sprint-4 (IB-28)"),
        ("R-011 (pgvector dim)", "Sprint-2 (IB-25) — debe hacerse antes de v0 ideally"),
        ("R-013 (brand voice)", "Sprint-3 (IB-34)"),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Riesgo"
    hdr[1].text = "Sprints / historias"
    for risk, sprints in risk_coverage:
        row = tbl.add_row().cells
        row[0].text = risk
        row[1].text = sprints

    path = HERE / "estrategia-tech.docx"
    doc.save(path)
    return path


def _add_story_table(doc, stories):
    """Helper: render a list of stories as a compact docx table."""
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for c, label in zip(hdr, ["ID", "Historia", "Pts", "Riesgo / Nota"]):
        c.text = label
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for s in stories:
        sid, title, points = s[0], s[2], s[7]
        risk, adr = s[12], s[13]
        row = tbl.add_row().cells
        row[0].text = sid
        row[1].text = title
        row[2].text = str(points)
        row[3].text = (risk if risk and risk != "—" else (adr if adr != "—" else ""))


# ============================================================
# 3. estrategia-exec.docx (Spanish, board / investor audience)
# ============================================================
def gen_docx_exec() -> Path:
    doc = Document()
    for s in doc.sections:
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)

    # ----- Portada -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Estrategia de Inversión Técnica")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("InkBytes — 6 semanas para pasar de MVP a producto cobrable")
    r.font.size = Pt(14)
    r.font.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Audiencia: Board / inversores / co-founder no técnico\n").italic = True
    meta.add_run("Versión 1.0 · 2026-06-02 · Pragmata Labs").italic = True

    doc.add_page_break()

    # ----- One-pager ejecutivo -----
    doc.add_heading("Resumen ejecutivo (1 página)", level=1)
    doc.add_paragraph(
        "InkBytes acaba de poner en producción su MVP: un lector de noticias pago, "
        "sin publicidad, donde cada evento se presenta como una página única "
        "consolidada desde múltiples fuentes. Las próximas 6 semanas convertirán esa "
        "demo invitada en un producto cobrable con calidad demostrable, controles "
        "operacionales serios y cumplimiento legal mínimo."
    )

    summary = doc.add_paragraph()
    summary.add_run("La inversión técnica propuesta: ").bold = True
    summary.add_run(
        "31 historias agrupadas en 8 áreas (seguridad, fiabilidad, datos, calidad "
        "LLM, deploy/CI, observabilidad, contratos, lector/monetización), "
        "ejecutadas en 3 sprints de 2 semanas. Costo en horas-ingeniero: ~95 "
        "puntos (≈ 4 semanas-persona). Costo en USD/mes nuevo: < $30 sobre el "
        "OPEX actual."
    )

    objetivo = doc.add_paragraph()
    objetivo.add_run("Objetivo de negocio: ").bold = True
    objetivo.add_run(
        "10+ usuarios pago activos con > 50% retention semanal al final de Sprint-4. "
        "Habilita la siguiente ronda de marketing porque el funnel completo (auth + "
        "billing + Stripe Checkout) está operativo, no es vaporware."
    )

    # ----- Tabla resumen del riesgo si NO se hace ----
    doc.add_heading("¿Por qué este orden importa? — Riesgo si no se ejecuta", level=1)
    risk_tbl = doc.add_table(rows=1, cols=3)
    risk_tbl.style = "Light Grid Accent 1"
    hdr = risk_tbl.rows[0].cells
    for c, label in zip(hdr, ["Si no se hace…", "Probabilidad", "Impacto en el negocio"]):
        c.text = label
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    risks = [
        ("Rotar tokens filtrados (R-009)", "Alta — el repo es público", "$2k-$50k en cargos no autorizados (DO Spaces, Anthropic, OpenAI). Pérdida de confianza si un atacante envenena la BD."),
        ("Cerrar hallucination del LLM (R-001)", "Alta — Haiku zero-shot tiene tasa de error real ~5-10%", "Reclamo legal de un outlet si reproducimos un hecho falso. Cancelación de suscripciones masivas si un usuario detecta una fabricación."),
        ("Backup restore drill (R-004)", "Baja por incidente, alta por agregación", "Pérdida de pages y sessions si Postgres muere. 8h de downtime para restaurar desde cero. Sin drill no sabemos si los backups funcionan."),
        ("Magic-link auth + Stripe (IB-49, IB-50)", "Crítica — sin esto no hay producto cobrable", "Reader queda en estado \"demo gratis\" indefinidamente. Cero ingresos."),
        ("Privacy policy + DSAR (IB-53)", "Media (sin penalización inmediata)", "Bloquea expansión en mercado europeo. Riesgo legal en DR ante cualquier reclamo de usuario."),
    ]
    for descr, prob, impact in risks:
        row = risk_tbl.add_row().cells
        row[0].text = descr
        row[1].text = prob
        row[2].text = impact

    doc.add_page_break()

    # ----- Visión del producto y diferenciación -----
    doc.add_heading("Visión del producto", level=1)
    doc.add_paragraph(
        "InkBytes es lo que sería Bloomberg si cobrara $9/mes en lugar de "
        "$2,000/mes y eliminara todo el ruido. Una página por evento. Sintetizada "
        "desde múltiples fuentes con IA. Con citas literales para verificar. "
        "Diseñada para quien necesita señal en LATAM bilingüe + tech/business global."
    )

    diff = doc.add_paragraph()
    diff.add_run("Diferenciación vs Ground News: ").bold = True
    diff.add_run(
        "Ground News etiqueta bias por outlet pero sigue siendo una lista de "
        "artículos. InkBytes consolida los artículos en una sola pieza de "
        "lectura, manteniendo la atribución. Más simple, más rápido de leer."
    )

    # ----- Inversión y retorno -----
    doc.add_heading("Inversión: 6 semanas, ~$30 OPEX adicional/mes", level=1)
    inv_tbl = doc.add_table(rows=1, cols=2)
    inv_tbl.style = "Light Grid Accent 1"
    for c, label in zip(inv_tbl.rows[0].cells, ["Concepto", "USD / mes"]):
        c.text = label
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    items_invest = [
        ("OPEX actual estimado (Droplet + Postgres + Spaces + Haiku)", "120"),
        ("Better Stack (logs) + Sentry (errors) — free tier", "0"),
        ("DO Managed Backup test (cuenta temporal)", "1"),
        ("CloudAMQP upgrade a paid (mTLS — opcional Sprint-3)", "19"),
        ("Stripe (sin costo fijo, fee por transacción)", "0"),
        ("Resend (magic-link emails, free tier ≤ 100/día)", "0"),
        ("Penetration test (one-time, Sprint-4)", "3,000-5,000 (one-time)"),
        ("Privacy policy template legal (one-time)", "500 (one-time)"),
        ("Staging Droplet (Sprint-3 opcional)", "24"),
        ("Total OPEX nuevo recurrente", "+$25-50/mes"),
        ("Total one-time", "$3,500-5,500"),
    ]
    for desc, usd in items_invest:
        row = inv_tbl.add_row().cells
        row[0].text = desc
        row[1].text = usd

    # ----- Métricas de éxito (board view) -----
    doc.add_heading("Cómo medirás el ROI al final de Sprint-4", level=1)
    metrics = [
        ("Comercial", "≥ 10 usuarios pago activos. ≥ $90/mes MRR. Retention semanal > 50%."),
        ("Operacional", "Uptime visible Reader ≥ 99.5% en últimas 4 semanas. Cost por artículo enriquecido ≤ $0.005. Cero incidentes P0 por secretos."),
        ("Calidad", "≥ 95% de pages publicadas pasan revisión humana sample. Cero reclamos formales de outlets por reproducción de hechos."),
        ("Cumplimiento", "Privacy policy publicada. DSAR endpoint funcional. Pentest completado sin findings Críticos."),
        ("Velocidad de iteración", "PR-to-prod < 1 hora. Mean time to recover de incidentes < 30 min."),
    ]
    for area, m in metrics:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{area}: ").bold = True
        p.add_run(m)

    # ----- Riesgos del plan -----
    doc.add_heading("Riesgos del plan (qué podría salir mal)", level=1)
    plan_risks = [
        ("Anthropic sube precios o Haiku se discontinúa", "Baja, alto impacto", "Curator está diseñado para swap LLM en una línea de config. Fallback a Gemini Flash documentado."),
        ("Outlets implementan paywalls duros / WAFs anti-scraping", "Media, medio impacto", "Diversificación: 12+ outlets ya configurados. Tradeoff: si > 30% caen, conversación con outlets sobre licensing."),
        ("Equipo de ingeniería pierde capacity (1 dev en lugar de 2)", "Media", "Sprint-4 es descopable (IB-24, IB-17, IB-28 son opcionales). Sprint-2 y Sprint-3 deben ejecutarse aunque sea con 1 dev."),
        ("Stripe rechaza la cuenta de DR", "Baja", "Plan B: Lemon Squeezy o Paddle (ambos soportan DR-based business)."),
        ("Ningún paying user en Sprint-4", "Media", "Indica problema de marketing o product-market fit, no técnico. La arquitectura permite pivotar el ICP sin reescribir el pipeline."),
    ]
    for descr, prob, mitig in plan_risks:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{descr} ({prob}): ").bold = True
        p.add_run(mitig)

    # ----- Cierre -----
    doc.add_heading("Decisión solicitada", level=1)
    doc.add_paragraph(
        "Aprobación del backlog y compromiso de capacity de 2 ingenieros por 6 "
        "semanas. El detalle ejecutable está en backlog.xlsx (priorización + plan "
        "por sprint) y las 31 historias individuales en docs/architecture/backlog/."
    )
    doc.add_paragraph(
        "El siguiente checkpoint es Sprint-2 review (T+14 días), donde validaremos "
        "que los 12 items críticos estén cerrados y que la base operacional "
        "soporte cobrar al primer usuario."
    )

    path = HERE / "estrategia-exec.docx"
    doc.save(path)
    return path


# ============================================================
# Main
# ============================================================
def main() -> None:
    xlsx = gen_xlsx()
    print(f"✓ {xlsx.name}  ({xlsx.stat().st_size // 1024} KB)")
    tech = gen_docx_tech()
    print(f"✓ {tech.name}  ({tech.stat().st_size // 1024} KB)")
    execu = gen_docx_exec()
    print(f"✓ {execu.name}  ({execu.stat().st_size // 1024} KB)")
    print(f"\nAll three written to: {HERE}")


if __name__ == "__main__":
    main()
